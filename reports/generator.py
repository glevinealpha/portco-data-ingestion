"""
AlphaFMC LP Report Generator.

Produces a Word (.docx) quarterly LP report using python-docx.
Claude API is used to generate narrative commentary.

Usage:
    python reports/generator.py --period "Q4 2024"
    python reports/generator.py --period "Q4 2024" --no-ai
"""

import argparse
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from pipeline.store import get_all_records, get_latest_records

# ── Brand colours ──────────────────────────────────────────────────────────────
NAVY_RGB  = RGBColor(0x0A, 0x16, 0x28)
GOLD_RGB  = RGBColor(0xC9, 0xA8, 0x4C)
GREY_RGB  = RGBColor(0x8A, 0x94, 0xA6)
WHITE_RGB = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_RGB = RGBColor(0xF5, 0xF6, 0xFA)


# ── Helpers ────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), str(val.get("sz", 4)))
            el.set(qn("w:color"), val.get("color", "000000"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def _para(doc, text: str, style: str = "Normal", bold: bool = False,
          italic: bool = False, size: int = 10, color: RGBColor = None,
          align=WD_ALIGN_PARAGRAPH.LEFT, space_before: int = 0,
          space_after: int = 4) -> None:
    p = doc.add_paragraph(style=style)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def _fmt(val, sym="", decimals=0) -> str:
    if val is None:
        return "—"
    return f"{sym}{val:,.{decimals}f}"


def _pct(val, decimals=1) -> str:
    if val is None:
        return "—"
    sign = "+" if val > 0 else ""
    return f"{sign}{val:.{decimals}f}%"


def _sym(currency: str) -> str:
    return "£" if currency == "GBP" else "€"


def rag_status(rec: dict) -> str:
    nd = rec.get("net_debt_ebitda", 0) or 0
    em = rec.get("vs_budget_ebitda_pct", 0) or 0
    if nd > 5 or em < -10:
        return "RED"
    if nd > 4 or em < -5:
        return "AMBER"
    return "GREEN"


def rag_color(status: str) -> str:
    return {"GREEN": "27AE60", "AMBER": "E67E22", "RED": "C0392B"}[status]


# ── AI narrative generation ────────────────────────────────────────────────────

def _generate_exec_summary(portco_data: list[dict], period: str) -> str:
    """Generate ~150-word executive summary via the configured AI provider."""
    from pipeline.ai_client import complete

    bullet_lines = []
    for r in portco_data:
        sym = _sym(r.get("currency", "GBP"))
        bullet_lines.append(
            f"- {r['portco_name']}: Revenue {_fmt(r['revenue'], sym)}"
            f"k, EBITDA {_fmt(r['ebitda'], sym)}k ({r.get('ebitda_margin', 0):.1f}% margin), "
            f"ND/EBITDA {r.get('net_debt_ebitda', 0):.2f}x, "
            f"EBITDA vs budget {_pct(r.get('vs_budget_ebitda_pct', 0))}"
        )

    prompt = (
        f"You are the Chief Investment Officer of AlphaFMC, a private equity firm. "
        f"Write a concise executive summary (~150 words) for the {period} quarterly LP report "
        f"covering the following portfolio company performance. "
        f"Use a professional, institutional tone. Highlight overall portfolio health, "
        f"key themes, and any items requiring LP attention.\n\n"
        + "\n".join(bullet_lines)
    )

    return complete(prompt, max_tokens=300)


def _generate_portco_commentary(rec: dict, prior: dict | None, period: str) -> str:
    """Generate ~100-word portco commentary via the configured AI provider."""
    from pipeline.ai_client import complete

    sym = _sym(rec.get("currency", "GBP"))
    curr_info = (
        f"Company: {rec['portco_name']}, {period}\n"
        f"Revenue: {_fmt(rec['revenue'], sym)}k "
        f"(vs budget: {_pct(rec.get('vs_budget_revenue_pct', 0))})\n"
        f"EBITDA: {_fmt(rec['ebitda'], sym)}k, margin {rec.get('ebitda_margin', 0):.1f}% "
        f"(vs budget: {_pct(rec.get('vs_budget_ebitda_pct', 0))})\n"
        f"Net Debt: {_fmt(rec['net_debt'], sym)}k, ND/EBITDA: "
        f"{rec.get('net_debt_ebitda', 0):.2f}x\n"
        f"Headcount: {rec.get('headcount', 0):,}"
    )

    prior_info = ""
    if prior:
        prior_info = (
            f"\nPrior quarter ({prior.get('period', 'prior')}):\n"
            f"Revenue: {_fmt(prior['revenue'], sym)}k, "
            f"EBITDA: {_fmt(prior['ebitda'], sym)}k, "
            f"Margin: {prior.get('ebitda_margin', 0):.1f}%"
        )

    prompt = (
        f"You are a portfolio manager at AlphaFMC writing an LP update. "
        f"Write a concise commentary (~100 words) on the following portfolio company's "
        f"quarterly performance. Highlight key developments vs prior quarter and vs budget. "
        f"Be factual and professional. Do not include headers.\n\n"
        f"{curr_info}{prior_info}"
    )

    return complete(prompt, max_tokens=180)


def _fallback_exec_summary(portco_data: list[dict], period: str) -> str:
    lines = [
        f"AlphaFMC's portfolio delivered mixed results in {period}. "
        "The five portfolio companies collectively demonstrated resilient top-line performance "
        "against a challenging macro backdrop, with most companies broadly in line with budget. "
    ]
    flagged = [r["portco_name"] for r in portco_data
               if rag_status(r) in ("RED", "AMBER")]
    if flagged:
        lines.append(
            f"{', '.join(flagged)} warrant close monitoring given leverage or budget variance metrics. "
        )
    lines.append(
        "Management teams are actively executing on operational improvement plans and the "
        "investment committee will continue to monitor performance against plan. "
        "Overall portfolio health remains satisfactory with positive revenue trends across the majority of holdings."
    )
    return " ".join(lines)


def _fallback_portco_commentary(rec: dict, prior: dict | None, period: str) -> str:
    sym = _sym(rec.get("currency", "GBP"))
    rev_var = rec.get("vs_budget_revenue_pct", 0) or 0
    ebit_var = rec.get("vs_budget_ebitda_pct", 0) or 0
    nd = rec.get("net_debt_ebitda", 0) or 0

    qual = "ahead of" if rev_var > 0 else "below"
    text = (
        f"{rec['portco_name']} reported revenue of {_fmt(rec['revenue'], sym)}k in {period}, "
        f"{abs(rev_var):.1f}% {qual} budget. EBITDA came in at {_fmt(rec['ebitda'], sym)}k "
        f"({rec.get('ebitda_margin', 0):.1f}% margin), "
        f"{'ahead of' if ebit_var > 0 else 'below'} budget by {abs(ebit_var):.1f}%. "
    )
    if prior:
        rev_qoq = (rec["revenue"] - prior["revenue"]) / prior["revenue"] * 100
        text += (
            f"Revenue grew {rev_qoq:.1f}% quarter-on-quarter. "
        )
    text += (
        f"Net debt / EBITDA stood at {nd:.2f}x. "
        f"Headcount was {rec.get('headcount', 0):,} at period end."
    )
    return text


# ── Document builder ───────────────────────────────────────────────────────────

def build_lp_report(period: str, use_ai: bool = True) -> Path:
    """
    Build and save the LP report Word document.
    Returns the path to the saved file.
    AI provider and credentials are read from config.properties.
    """

    all_records = get_all_records()
    if not all_records:
        raise RuntimeError("No data in database. Run the data generator first.")

    # Filter to requested period
    period_records = [r for r in all_records if r["period"] == period]
    if not period_records:
        available = sorted(set(r["period"] for r in all_records))
        raise ValueError(f"Period '{period}' not found. Available: {available}")

    # Build prior-quarter lookup
    quarter_order = {f"Q{q} {y}": i for i, (q, y) in enumerate(
        [(q, yr) for yr in [2023, 2024] for q in [1, 2, 3, 4]]
    )}
    prior_lookup: dict[str, dict] = {}
    for rec in period_records:
        portco_id = rec["portco_id"]
        curr_idx = quarter_order.get(period, -1)
        if curr_idx > 0:
            for prior_period, idx in sorted(quarter_order.items(), key=lambda x: x[1], reverse=True):
                if idx < curr_idx:
                    candidates = [r for r in all_records
                                  if r["portco_id"] == portco_id
                                  and r["period"] == prior_period]
                    if candidates:
                        prior_lookup[portco_id] = candidates[0]
                        break

    doc = Document()

    # ── Document-wide margins ──────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Cover page ─────────────────────────────────────────────────────────────
    doc.add_paragraph()  # top spacer

    # AlphaFMC title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("AlphaFMC")
    title_run.font.size = Pt(36)
    title_run.font.bold = True
    title_run.font.color.rgb = NAVY_RGB

    # Gold rule
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule_run = rule.add_run("─" * 40)
    rule_run.font.color.rgb = GOLD_RGB
    rule_run.font.size = Pt(12)

    # Report type
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("QUARTERLY LP REPORT")
    sub_run.font.size = Pt(18)
    sub_run.font.bold = True
    sub_run.font.color.rgb = GREY_RGB
    sub_run.font.letter_spacing = Pt(2)

    doc.add_paragraph()

    # Period
    per_para = doc.add_paragraph()
    per_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    per_run = per_para.add_run(period)
    per_run.font.size = Pt(28)
    per_run.font.bold = True
    per_run.font.color.rgb = GOLD_RGB

    doc.add_paragraph()

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Prepared: {datetime.now().strftime('%d %B %Y')}")
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = GREY_RGB

    doc.add_paragraph()
    doc.add_paragraph()

    # Confidentiality
    conf_para = doc.add_paragraph()
    conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf_run = conf_para.add_run("CONFIDENTIAL — NOT FOR DISTRIBUTION")
    conf_run.font.size = Pt(9)
    conf_run.font.bold = True
    conf_run.font.color.rgb = GREY_RGB

    doc.add_page_break()

    # ── Executive Summary ──────────────────────────────────────────────────────
    h_para = doc.add_paragraph()
    h_run = h_para.add_run("Executive Summary")
    h_run.font.size = Pt(16)
    h_run.font.bold = True
    h_run.font.color.rgb = NAVY_RGB
    h_para.paragraph_format.space_after = Pt(2)

    # Gold underline
    hr = doc.add_paragraph()
    hr_run = hr.add_run("─" * 80)
    hr_run.font.color.rgb = GOLD_RGB
    hr_run.font.size = Pt(8)
    hr.paragraph_format.space_after = Pt(8)

    print(f"Generating executive summary {'(AI)' if use_ai else '(template)'}...")
    if use_ai:
        try:
            exec_text = _generate_exec_summary(period_records, period)
        except Exception as e:
            print(f"  AI error: {e}. Using template.")
            exec_text = _fallback_exec_summary(period_records, period)
    else:
        exec_text = _fallback_exec_summary(period_records, period)

    exec_para = doc.add_paragraph(exec_text)
    exec_para.paragraph_format.space_after = Pt(12)
    for run in exec_para.runs:
        run.font.size = Pt(10)

    # ── Portfolio Summary Table ────────────────────────────────────────────────
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("Portfolio Summary")
    h2_run.font.size = Pt(13)
    h2_run.font.bold = True
    h2_run.font.color.rgb = NAVY_RGB

    headers = ["Company", "Sector", "Currency", "Revenue", "EBITDA",
               "Margin", "ND/EBITDA", "Rev vs Bgt", "EBITDA vs Bgt", "Status"]
    tbl = doc.add_table(rows=1 + len(period_records), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = WHITE_RGB
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, "0A1628")

    # Data rows
    for row_idx, rec in enumerate(period_records):
        row = tbl.rows[row_idx + 1]
        sym = _sym(rec.get("currency", "GBP"))
        status = rag_status(rec)
        row_data = [
            rec["portco_name"],
            rec.get("sector", ""),
            rec.get("currency", ""),
            _fmt(rec["revenue"], sym),
            _fmt(rec["ebitda"], sym),
            f'{rec.get("ebitda_margin", 0):.1f}%',
            f'{rec.get("net_debt_ebitda", 0):.2f}x',
            _pct(rec.get("vs_budget_revenue_pct", 0)),
            _pct(rec.get("vs_budget_ebitda_pct", 0)),
            status,
        ]
        bg = "F5F6FA" if row_idx % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            align = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].alignment = align
            _set_cell_bg(cell, bg)
            # Colour status cell
            if ci == len(headers) - 1:
                _set_cell_bg(cell, rag_color(status))
                cell.paragraphs[0].runs[0].font.color.rgb = WHITE_RGB
                cell.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()
    doc.add_page_break()

    # ── Per-portco pages ───────────────────────────────────────────────────────
    for rec in period_records:
        sym = _sym(rec.get("currency", "GBP"))
        prior = prior_lookup.get(rec["portco_id"])
        status = rag_status(rec)

        # Section header
        co_h = doc.add_paragraph()
        co_run = co_h.add_run(rec["portco_name"])
        co_run.font.size = Pt(15)
        co_run.font.bold = True
        co_run.font.color.rgb = NAVY_RGB

        sub_line = doc.add_paragraph()
        sub_run = sub_line.add_run(
            f"{rec.get('sector', '')}  |  {rec.get('geography', '')}  |  "
            f"{rec.get('stage', '')}  |  {period}"
        )
        sub_run.font.size = Pt(9)
        sub_run.font.color.rgb = GREY_RGB
        sub_line.paragraph_format.space_after = Pt(2)

        hr2 = doc.add_paragraph()
        hr2_run = hr2.add_run("─" * 80)
        hr2_run.font.color.rgb = GOLD_RGB
        hr2_run.font.size = Pt(8)
        hr2.paragraph_format.space_after = Pt(8)

        # KPI table
        kpi_rows = [
            ("Revenue",          _fmt(rec["revenue"], sym), f"{sym}000s"),
            ("Gross Profit",     _fmt(rec["gross_profit"], sym), f"{sym}000s"),
            ("EBITDA",           _fmt(rec["ebitda"], sym), f"{sym}000s"),
            ("EBITDA Margin",    f'{rec.get("ebitda_margin", 0):.1f}%', ""),
            ("EBIT",             _fmt(rec["ebit"], sym), f"{sym}000s"),
            ("Net Income",       _fmt(rec["net_income"], sym), f"{sym}000s"),
            ("Net Debt",         _fmt(rec["net_debt"], sym), f"{sym}000s"),
            ("Net Debt / EBITDA",f'{rec.get("net_debt_ebitda", 0):.2f}x', ""),
            ("Headcount",        f'{rec.get("headcount", 0):,}', "FTEs"),
            ("Rev vs Budget",    _pct(rec.get("vs_budget_revenue_pct", 0)), ""),
            ("EBITDA vs Budget", _pct(rec.get("vs_budget_ebitda_pct", 0)), ""),
        ]

        kt = doc.add_table(rows=1 + len(kpi_rows), cols=3)
        kt.style = "Table Grid"

        # KPI table header
        for ci, h in enumerate(["Metric", "Value", "Unit"]):
            cell = kt.rows[0].cells[ci]
            cell.text = h
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = WHITE_RGB
            cell.paragraphs[0].runs[0].font.size = Pt(8)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_bg(cell, "0A1628")

        for ri, (metric, value, unit) in enumerate(kpi_rows):
            row = kt.rows[ri + 1]
            for ci, val in enumerate([metric, value, unit]):
                cell = row.cells[ci]
                cell.text = val
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                cell.paragraphs[0].alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER
                )
                bg = "F5F6FA" if ri % 2 == 0 else "FFFFFF"
                _set_cell_bg(cell, bg)
                # Bold totals
                if metric in ("Revenue", "EBITDA", "Net Debt"):
                    cell.paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()

        # Status
        status_para = doc.add_paragraph()
        status_run = status_para.add_run(f"RAG Status: {status}")
        status_run.font.bold = True
        status_run.font.size = Pt(10)
        status_run.font.color.rgb = RGBColor.from_string(rag_color(status))
        status_para.paragraph_format.space_after = Pt(8)

        # Commentary
        comm_h = doc.add_paragraph()
        comm_run = comm_h.add_run("Management Commentary")
        comm_run.font.size = Pt(11)
        comm_run.font.bold = True
        comm_run.font.color.rgb = NAVY_RGB
        comm_h.paragraph_format.space_after = Pt(4)

        print(f"Generating commentary for {rec['portco_name']} {'(AI)' if use_ai else '(template)'}...")
        if use_ai:
            try:
                commentary = _generate_portco_commentary(rec, prior, period)
            except Exception as e:
                print(f"  AI error: {e}. Using template.")
                commentary = _fallback_portco_commentary(rec, prior, period)
        else:
            commentary = _fallback_portco_commentary(rec, prior, period)

        comm_para = doc.add_paragraph(commentary)
        for run in comm_para.runs:
            run.font.size = Pt(10)
        comm_para.paragraph_format.space_after = Pt(6)

        doc.add_page_break()

    # ── Save ───────────────────────────────────────────────────────────────────
    reports_dir = PROJECT_ROOT / "reports" / "output"
    reports_dir.mkdir(parents=True, exist_ok=True)

    safe_period = period.replace(" ", "_")
    out_path = reports_dir / f"AlphaFMC_LP_Report_{safe_period}.docx"
    doc.save(str(out_path))
    return out_path


# ── CLI ────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="AlphaFMC LP Report Generator")
    parser.add_argument("--period", default="Q4 2024",
                        help="Reporting period, e.g. 'Q4 2024'")
    parser.add_argument("--no-ai", action="store_true",
                        help="Skip Claude API calls, use template text")
    args = parser.parse_args()

    try:
        out = build_lp_report(
            period=args.period,
            use_ai=not args.no_ai,
        )
        print(f"\nLP Report saved: {out}")
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)
